#!/usr/bin/env python3
"""生成南京禄通生命科学有限公司 新员工转正申请表 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── 辅助样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_centered_text(cell, text, bold=False, size=Pt(12), font_name='宋体'):
    """在单元格中添加居中文字"""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_left_text(cell, text, bold=False, size=Pt(12), font_name='宋体'):
    """在单元格中添加左对齐文字"""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def merge_and_set(cells, text, bold=True, size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """合并单元格并设置文字"""
    merged = cells[0].merge(cells[-1])
    merged.paragraphs[0].alignment = alignment
    run = merged.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return merged

# ========================================================================
# 第一页：新员工转正申请表（第一部分：员工个人述职报告）
# ========================================================================

# ── 公司名称 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('南京禄通生命科学有限公司')
run.bold = True
run.font.size = Pt(18)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.space_after = Pt(4)

# ── 表格标题 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('新员工转正申请表')
run.bold = True
run.font.size = Pt(16)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.space_after = Pt(8)

# ── 第一部分标题 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('第一部分：员工个人述职报告')
run.bold = True
run.font.size = Pt(14)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.space_after = Pt(6)

# ── 员工信息表 ──
info_table = doc.add_table(rows=2, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

# Row 0: 姓名 / 刘婷婷 / 部门 / 运营部
set_cell_shading(info_table.cell(0, 0), 'D9E2F3')
add_centered_text(info_table.cell(0, 0), '姓名', bold=True)
add_centered_text(info_table.cell(0, 1), '刘婷婷')
set_cell_shading(info_table.cell(0, 2), 'D9E2F3')
add_centered_text(info_table.cell(0, 2), '部门', bold=True)
add_centered_text(info_table.cell(0, 3), '运营部')

# Row 1: 职位 / 专员 / 试用期时间 / 2016.3.10 - 2016.6.9
set_cell_shading(info_table.cell(1, 0), 'D9E2F3')
add_centered_text(info_table.cell(1, 0), '职位', bold=True)
add_centered_text(info_table.cell(1, 1), '专员')
set_cell_shading(info_table.cell(1, 2), 'D9E2F3')
add_centered_text(info_table.cell(1, 2), '试用期时间', bold=True)
add_centered_text(info_table.cell(1, 3), '2016.3.10 - 2016.6.9')

doc.add_paragraph()  # 空行

# ── 述职要求 ──
p = doc.add_paragraph()
run = p.add_run('个人述职报告要求：')
run.bold = True
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for req in ['1. 主要工作任务', '2. 完成情况', '3. 存在的问题及后期计划']:
    p = doc.add_paragraph(req)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# ── 述职内容表 ──
content_table = doc.add_table(rows=1, cols=1)
content_table.alignment = WD_TABLE_ALIGNMENT.CENTER
content_table.style = 'Table Grid'

cell = content_table.cell(0, 0)

# 1. 主要工作任务
p = cell.add_paragraph()
run = p.add_run('1. 主要工作任务：')
run.bold = True
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

items_1 = [
    '① 单据制作：用ERP完成采购、销售合同、出库单、采购发票、付款申请、销售发票的系统记录。',
    '② 订单跟进：与工厂、客户签订合同；发货前核对COA、标签以及客户要求；核对发货明细、送货单、收货单；跟进开票资料。',
    '③ 资料归档：合同、COA、标签、送货单、物流信息、收货确认单、发票分类存档。'
]
for item in items_1:
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 2. 完成情况
p = cell.add_paragraph()
run = p.add_run('2. 完成情况：')
run.bold = True
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = cell.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
run = p.add_run('及时响应业务、客户、财务的单据查询；妥善处理单据异常问题；未因单证问题产生客诉、财务纠纷。')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 3. 存在的问题及后期计划
p = cell.add_paragraph()
run = p.add_run('3. 存在的问题及后期计划：')
run.bold = True
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = cell.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
run = p.add_run('多订单集中出货高峰期，制单节奏紧张，偶尔出现信息疏漏，需二次复核修正。未来还需吃透特殊业务流程，对于特殊订单，要做前置风险核查。')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 删除第一个空段落
if cell.paragraphs[0].text == '':
    p_elem = cell.paragraphs[0]._element
    p_elem.getparent().remove(p_elem)

doc.add_paragraph()  # 空行

# ── 签名与日期 ──
sign_table = doc.add_table(rows=1, cols=4)
sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sign_table.style = 'Table Grid'

# 签名栏
set_cell_shading(sign_table.cell(0, 0), 'D9E2F3')
add_centered_text(sign_table.cell(0, 0), '员工签名', bold=True)
add_centered_text(sign_table.cell(0, 1), '刘婷婷')

# 日期栏
set_cell_shading(sign_table.cell(0, 2), 'D9E2F3')
add_centered_text(sign_table.cell(0, 2), '日期', bold=True)
add_centered_text(sign_table.cell(0, 3), '2026.6.9')

# ── 页码 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('第1页 共2页')
run.font.size = Pt(10)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========================================================================
# 分页：第二页 - 第二部分：转正审核表
# ========================================================================
doc.add_page_break()

# ── 公司名称 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('南京禄通生命科学有限公司')
run.bold = True
run.font.size = Pt(18)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.space_after = Pt(4)

# ── 表格标题 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('第二部分：转正审核表')
run.bold = True
run.font.size = Pt(16)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.space_after = Pt(8)

# ── 转正审核表格 ──
audit_table = doc.add_table(rows=4, cols=2)
audit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
audit_table.style = 'Table Grid'

# 列宽设置
for row in audit_table.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(12)

# Row 0: 直属上级意见
row0 = audit_table.rows[0]
row0.height = Cm(4.5)
set_cell_shading(row0.cells[0], 'D9E2F3')
add_centered_text(row0.cells[0], '直属上级意见', bold=True)

opinion1 = (
    '该员工试用期上进心足、乐于协作、基础工作可胜任，'
    '少量细节需优化指导，整体符合岗位转正标准，同意按期转正。'
)
p = row0.cells[1].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(opinion1)
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p2 = row0.cells[1].add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p2.add_run('签名：吴丹    日期：2026.06.10')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Row 1: 部门负责人意见
row1 = audit_table.rows[1]
row1.height = Cm(2.5)
set_cell_shading(row1.cells[0], 'D9E2F3')
add_centered_text(row1.cells[0], '部门负责人意见', bold=True)

p = row1.cells[1].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('同意')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p2 = row1.cells[1].add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p2.add_run('签名：吴丹    日期：2026.06.10')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Row 2: 综合管理部意见
row2 = audit_table.rows[2]
row2.height = Cm(2.5)
set_cell_shading(row2.cells[0], 'D9E2F3')
add_centered_text(row2.cells[0], '综合管理部意见', bold=True)

p = row2.cells[1].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('同意')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p2 = row2.cells[1].add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p2.add_run('签名：柳宏彦    日期：2026.6.10')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Row 3: 总经理意见
row3 = audit_table.rows[3]
row3.height = Cm(2.5)
set_cell_shading(row3.cells[0], 'D9E2F3')
add_centered_text(row3.cells[0], '总经理意见', bold=True)

p = row3.cells[1].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('同意')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p2 = row3.cells[1].add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p2.add_run('签名：______    日期：6/10')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 页码 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('第2页 共2页')
run.font.size = Pt(10)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 保存 ──
output_path = os.path.expanduser('~/南京禄通生命科学有限公司_新员工转正申请表.docx')
doc.save(output_path)
print(f'文档已保存至: {output_path}')
